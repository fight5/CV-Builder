"""Pipeline OPTIMUM — CV + Lettre via DeepSeek/Gemini, sans JSON.

Préférences utilisateur :
- template       : "optimum" | "minimal"
- language       : "Français" | "English"
- accent_hex     : couleur d'accent (hex)
- leftbg_hex     : couleur de fond du bandeau gauche (optimum uniquement)
- include_photo  : bool
- photo_path     : str | None  (chemin local accessible à pdflatex)
- aggressive     : bool  (laisser le LLM lister des compétences/outils
                         non explicites dans le CV source — assumé par
                         le candidat)
- company        : str   (facultatif — utilisé pour nommer les fichiers)

Sortie : dict avec cv_latex, cv_pdf_bytes, letter_body, letter_latex,
letter_pdf_bytes, errors, et les noms de fichiers suggérés.
"""

from __future__ import annotations

import json
import logging
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

PROJECT_ROOT = Path(__file__).resolve().parent.parent
TEMPLATES_DIR = PROJECT_ROOT / "templates"
ASSETS_DIR = TEMPLATES_DIR / "assets"   # logos, photo, images utilisateur
TEMPLATE_LETTER = TEMPLATES_DIR / "letter.tex"
OUTPUT_DIR = PROJECT_ROOT / "outputs"

_MONTHS_FR = [
    "janvier", "février", "mars", "avril", "mai", "juin",
    "juillet", "août", "septembre", "octobre", "novembre", "décembre",
]
_MONTHS_EN = [
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December",
]


# ── Préférences ──────────────────────────────────────────────────────────────
@dataclass
class CVPreferences:
    template: str = "optimum"          # optimum | minimal
    language: str = "Français"         # Français | English
    accent_hex: str = "#006699"        # bleu pétrole par défaut
    leftbg_hex: str = "#172E4A"        # bleu marine bandeau gauche
    include_photo: bool = False
    photo_path: Optional[str] = None
    # QR code (optionnel) — chemin vers l'image, label affiché en dessous
    qr_code_path: Optional[str] = None
    qr_code_label: str = ""
    # Photo dans le bandeau gauche (optionnel — distinct de la photo header)
    sidebar_photo_path: Optional[str] = None
    aggressive: bool = True            # le candidat assume les ajouts
    company: str = ""                  # entreprise (pour nommage + lettre)
    generate_letter: bool = False      # générer la lettre de motivation
    # Fichiers extra uploadés par l'utilisateur : {nom_cible.ext: chemin_source}
    # Copiés dans outputs/ avant compilation (logos, photo supplémentaire…)
    extra_assets: dict = field(default_factory=dict)


# ── Extraction texte ─────────────────────────────────────────────────────────
def extract_text_from_pdf(file_path: str) -> str:
    import fitz
    doc = fitz.open(file_path)
    try:
        return "\n".join(page.get_text("text") for page in doc)
    finally:
        doc.close()


def extract_text_from_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


def extract_cv_text(uploaded_file) -> str:
    """Streamlit UploadedFile -> texte brut. PDF/DOCX."""
    suffix = Path(uploaded_file.name).suffix.lower()
    data = uploaded_file.read()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(data)
        tmp_path = tmp.name
    try:
        if suffix == ".pdf":
            return extract_text_from_pdf(tmp_path)
        if suffix in (".docx", ".doc"):
            return extract_text_from_docx(tmp_path)
        return data.decode("utf-8", errors="ignore")
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


# ── LLM (DeepSeek prioritaire, Gemini en fallback) ───────────────────────────
def _provider() -> str:
    explicit = (os.getenv("LLM_PROVIDER") or "").strip().lower()
    if explicit in {"deepseek", "gemini"}:
        return explicit
    if os.getenv("DEEPSEEK_API_KEY"):
        return "deepseek"
    return "gemini"


def _call_deepseek(prompt: str, temperature: float) -> str:
    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        raise RuntimeError("DEEPSEEK_API_KEY manquante.")
    try:
        from openai import OpenAI
    except ImportError as e:
        raise RuntimeError("Installer : pip install openai") from e
    client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
    model = os.getenv("DEEPSEEK_MODEL", "deepseek-chat")
    resp = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=temperature,
        max_tokens=8192,
    )
    text = (resp.choices[0].message.content or "").strip()
    if not text:
        raise RuntimeError("Réponse DeepSeek vide.")
    return text


def _call_gemini(prompt: str, temperature: float) -> str:
    api_key = os.getenv("GOOGLE_API_KEY") or os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("GOOGLE_API_KEY manquante.")
    from langchain_google_genai import ChatGoogleGenerativeAI
    from langchain_core.messages import HumanMessage
    llm = ChatGoogleGenerativeAI(
        model=os.getenv("GEMINI_MODEL", "gemini-2.5-flash"),
        google_api_key=api_key,
        temperature=temperature,
        max_output_tokens=8192,
    )
    resp = llm.invoke([HumanMessage(content=prompt)])
    text = (resp.content or "").strip()
    if not text:
        raise RuntimeError("Réponse Gemini vide.")
    return text


def generate_with_llm(prompt: str, temperature: float = 0.35) -> str:
    provider = _provider()
    if provider == "deepseek":
        try:
            return _call_deepseek(prompt, temperature)
        except Exception as e:
            logger.warning("DeepSeek KO (%s) — fallback Gemini.", e)
            return _call_gemini(prompt, temperature)
    return _call_gemini(prompt, temperature)


# Alias historique.
generate_with_gemini = generate_with_llm


# ── Prompt JSON structuré ────────────────────────────────────────────────────
# Le LLM ne renvoie QUE les données (JSON), jamais le LaTeX complet.
# La structure du template est préservée à 100% — seul le contenu change.
_PROMPT_STRUCTURED_JSON = """\
Tu es expert en recrutement, ATS, et rédaction de CV haut de gamme.

Ta mission : extraire et adapter les informations du CV source pour qu'elles correspondent
parfaitement à l'offre d'emploi, puis les retourner sous forme JSON.

RÈGLES STRICTES :
1. Ne JAMAIS inventer une expérience, une entreprise, des dates ou un diplôme.
2. Ne JAMAIS inventer une certification absente du CV source.
3. Conserver les noms d'entreprises et d'écoles tels quels.
4. header_title = intitulé EXACT du poste tel qu'écrit dans l'offre.
5. skills / tools : inclure TOUTES les technologies / concepts de l'offre
   même s'ils n'apparaissent pas dans le CV source (le candidat assume).
6. Enrichir les bullets d'expérience avec les outils/résultats de l'offre
   tout en restant cohérent avec la mission réelle.
7. Quantifier le plus possible (%, volumes, échelle).
8. summary : liste de 3 bullet points PERCUTANTS (pas de prose, pas de paragraphe).
   Chaque bullet : 1 ligne, impact chiffré ou compétence clé, orienté offre. Court, frappant, précis.
   Format JSON attendu : liste de strings, ex. ["Bullet 1.", "Bullet 2.", "Bullet 3."]
9. Si un champ liste est vide (ex. aucune certification), mettre [].
10. IMPORTANT — retourner du texte BRUT dans les valeurs JSON.
    Ne pas utiliser de caractères d'échappement LaTeX (pas de \\& \\% \\# \\_).
    Le code Python se chargera de l'échappement LaTeX après parsing.

OBJECTIF REMPLISSAGE PAGE A4 COMPLÈTE :
Le CV DOIT remplir intégralement la page A4 — ni débordement ni espace vide.
Génère un contenu RICHE, COMPLET et PERCUTANT :
- summary : 3-4 lignes dynamiques, orientées offre, avec chiffres/impact.
- experiences : TOUTES les expériences pertinentes du CV source (jusqu'à 5).
  → 4-5 bullets forts par expérience récente, 3 bullets pour les plus anciennes.
  → Chaque bullet : 1 ligne avec impact mesurable (%, volumes, gains de temps…).
  → Enrichir avec les outils/méthodologies de l'offre, cohérents avec le poste réel.
- education : TOUTES les formations du CV source.
- skills : les compétences LES PLUS PERTINENTES pour le poste (6-8 items MAX — qualité > quantité).
- tools : les outils LES PLUS PERTINENTS de l'offre + du CV (6-8 items MAX).
- certifications : TOUTES les certifications du CV source (≤ 4).
- qualities : 4-5 items MAX (soft skills distinctifs, pas génériques).
- languages : TOUTES les langues maîtrisées avec le niveau.
- awards : distinctions du CV source (≤ 3).
- interests : 2-3 items MAX (pertinents au secteur).

LANGUE de tout le contenu : {language_label}

OFFRE D'EMPLOI :
{job_offer}

CV SOURCE :
{source_cv}

Retourne UNIQUEMENT le JSON suivant, complété. Aucun autre texte avant ou après.

{{
  "full_name": "",
  "header_title": "",
  "email": "",
  "linkedin_url": "",
  "linkedin_handle": "",
  "phone": "",
  "location": "",
  "summary": ["Bullet 1 percutant avec chiffre.", "Bullet 2 compétence clé.", "Bullet 3 valeur ajoutée."],
  "skills": [],
  "tools": [],
  "languages": [],
  "certifications": [],
  "qualities": [],
  "interests": [],
  "awards": [],
  "experiences": [
    {{
      "title": "",
      "company": "",
      "dates": "",
      "bullets": ["Réalisation quantifiée 1.", "Réalisation 2."]
    }}
  ],
  "education": [
    {{
      "degree": "",
      "school": "",
      "dates": ""
    }}
  ]
}}
"""

_PROMPT_LETTER = """\
Tu es un expert en rédaction de lettres de motivation haut de gamme.

Rédige une lettre de motivation personnalisée et convaincante en te basant sur :
- l'offre d'emploi,
- le CV optimisé du candidat.

OBJECTIFS :
1. Montrer clairement l'adéquation entre le candidat et le poste.
2. Mettre en avant les expériences les plus pertinentes.
3. Adopter un ton professionnel et naturel.
4. Faire ressortir la valeur ajoutée du candidat.
5. Donner envie au recruteur de le rencontrer.

CONTRAINTES :
- Une page maximum.
- {language_label} professionnel.
- Pas de phrases génériques.
- Pas de commentaire.
- Ne reprends pas le CV : valorise ce qui n'est pas explicite (motivation, vision, projection).
- 3 paragraphes séparés par une ligne vide.

OFFRE :
{job_offer}

CV OPTIMISÉ :
{optimized_cv}

Retourne uniquement le corps de la lettre, sans salutation d'ouverture, sans formule de fermeture, sans en-tête, sans date, sans signature.
"""


# ── Helpers LaTeX ────────────────────────────────────────────────────────────
def _hex_to_rgb(hex_color: str) -> str:
    """`#006699` -> `0,102,153`. Tolère sans #."""
    h = (hex_color or "").lstrip("#").strip()
    if len(h) != 6:
        return "0,102,153"
    try:
        return f"{int(h[0:2], 16)},{int(h[2:4], 16)},{int(h[4:6], 16)}"
    except ValueError:
        return "0,102,153"


def _latex_escape(s: str) -> str:
    if not isinstance(s, str):
        s = str(s)
    repl = [
        ("\\", r"\textbackslash{}"),
        ("&", r"\&"),
        ("%", r"\%"),
        ("$", r"\$"),
        ("#", r"\#"),
        ("_", r"\_"),
        ("{", r"\{"),
        ("}", r"\}"),
        ("~", r"\textasciitilde{}"),
        ("^", r"\textasciicircum{}"),
    ]
    for k, v in repl:
        s = s.replace(k, v)
    return s


def _french_date() -> str:
    d = datetime.now()
    return f"{d.day} {_MONTHS_FR[d.month - 1]} {d.year}"


def _english_date() -> str:
    d = datetime.now()
    return f"{_MONTHS_EN[d.month - 1]} {d.day}, {d.year}"


def _paragraphs_to_latex(body: str) -> str:
    body = (body or "").strip()
    if not body:
        return ""
    paragraphs = re.split(r"\n\s*\n", body)
    cleaned = []
    for p in paragraphs:
        line = " ".join(seg.strip() for seg in p.splitlines() if seg.strip())
        cleaned.append(_latex_escape(line))
    return r" \\[0.4em] ".join(cleaned)


# ── Assets (logos, photo) ────────────────────────────────────────────────────

def _copy_assets(
    output_dir: Path,
    extra_assets: Optional[dict] = None,
    prefs: Optional["CVPreferences"] = None,
) -> None:
    """Copie tous les fichiers image de templates/assets/ vers output_dir.

    Appelé avant chaque compilation pdflatex pour que les \\includegraphics
    et \\safeimage puissent trouver les fichiers.
    Copie aussi les extra_assets uploadés par l'utilisateur : {nom_cible: chemin_source}.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    if ASSETS_DIR.exists():
        for f in ASSETS_DIR.iterdir():
            if f.suffix.lower() in {".png", ".jpg", ".jpeg", ".eps", ".pdf"}:
                dest = output_dir / f.name
                try:
                    shutil.copyfile(f, dest)
                except OSError as e:
                    logger.debug("Asset copy skipped (%s): %s", f.name, e)
    # Fichiers uploadés par l'utilisateur (logos d'entreprise, photo…)
    if extra_assets:
        for target_name, src_path in extra_assets.items():
            if src_path and Path(src_path).exists():
                dest = output_dir / target_name
                try:
                    shutil.copyfile(src_path, dest)
                    logger.debug("Extra asset copié : %s → %s", src_path, dest)
                except OSError as e:
                    logger.debug("Extra asset non copié (%s): %s", target_name, e)
    # QR code
    if prefs and prefs.qr_code_path and Path(prefs.qr_code_path).exists():
        try:
            shutil.copyfile(prefs.qr_code_path, output_dir / "qrcode.png")
        except OSError as e:
            logger.debug("QR code non copié : %s", e)
    # Photo dans le bandeau gauche (sidebar)
    if prefs and prefs.sidebar_photo_path and Path(prefs.sidebar_photo_path).exists():
        try:
            shutil.copyfile(prefs.sidebar_photo_path, output_dir / "sidebar_photo.png")
        except OSError as e:
            logger.debug("Sidebar photo non copiée : %s", e)


def _find_logo(name: str, extra_filenames: Optional[dict] = None) -> Optional[str]:
    """Cherche un logo correspondant au nom donné.

    Cherche d'abord dans les extra_filenames (uploadés par l'utilisateur),
    puis dans templates/assets/.
    Retourne le nom de fichier (pas le chemin complet) ou None.
    Exemples : "Thales Alenia Space" → "Thales.png",
               "Sanofi" → "Sanofi.png",
               "EPF Montpellier" → "EPF.png".
    """
    if not name:
        return None
    name_lower = name.lower()

    def _score(stem: str) -> int:
        stem = stem.lower()
        if stem in name_lower or name_lower in stem:
            return len(stem)
        parts = [p for p in stem.split() if len(p) > 2]
        matching = [p for p in parts if p in name_lower]
        return max((len(p) for p in matching), default=0)

    best: Optional[str] = None
    best_score = 0

    # 1. Chercher dans les fichiers uploadés par l'utilisateur
    if extra_filenames:
        for fname in extra_filenames:
            if Path(fname).suffix.lower() not in {".png", ".jpg", ".jpeg"}:
                continue
            s = _score(Path(fname).stem)
            if s > best_score:
                best = fname
                best_score = s

    # 2. Chercher dans templates/assets/
    if ASSETS_DIR.exists():
        for f in ASSETS_DIR.iterdir():
            if f.suffix.lower() not in {".png", ".jpg", ".jpeg"}:
                continue
            s = _score(f.stem)
            if s > best_score:
                best = f.name
                best_score = s

    return best


# ── Build photo block ────────────────────────────────────────────────────────
def _photo_block(prefs: CVPreferences) -> str:
    """Bloc minipage avec photo, ou minipage vide. Garde l'alignement du header.

    Priorités :
    1. Photo uploadée via l'UI (prefs.photo_path) si include_photo=True.
    2. Auto-détection depuis templates/assets/ (photo_didentite.png ou photo.png/jpg).
    3. Bloc vide (0pt width).
    """
    photo_src: Optional[str] = None

    # Priorité 1 : photo explicitement fournie via UI
    if prefs.include_photo and prefs.photo_path:
        photo_src = prefs.photo_path

    # Priorité 2 : auto-détection dans templates/assets/
    if not photo_src:
        for candidate_name in ("photo_didentite.png", "photo_didentite.PNG",
                               "photo.png", "photo.jpg", "profile.png"):
            candidate = ASSETS_DIR / candidate_name
            if candidate.exists():
                photo_src = str(candidate)
                break

    if photo_src:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        dest = OUTPUT_DIR / "photo_didentite.png"
        try:
            shutil.copyfile(photo_src, dest)
            return (
                "\\begin{minipage}[t]{0.23\\textwidth}\n"
                "\\vspace{0pt}\n"
                "\\includegraphics[width=\\linewidth]{photo_didentite.png}\n"
                "\\end{minipage}\n\\hfill"
            )
        except OSError as e:
            logger.warning("Photo non copiable : %s", e)
    return (
        "\\begin{minipage}[t]{0.0\\textwidth}\\vspace{0pt}\\end{minipage}\\hfill"
    )


def _photo_block_minimal(prefs: CVPreferences) -> str:
    """Photo centrée pour le template minimal."""
    if prefs.include_photo and prefs.photo_path:
        try:
            OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
            dest = OUTPUT_DIR / "photo_didentite.png"
            shutil.copyfile(prefs.photo_path, dest)
            return (
                "\\begin{center}\n"
                "\\includegraphics[width=3.2cm,height=3.2cm,keepaspectratio]{photo_didentite.png}\n"
                "\\end{center}"
            )
        except OSError:
            pass
    return ""


def _qr_code_block(prefs: CVPreferences) -> str:
    """Bloc QR code pour le bandeau gauche.

    Affiche l'image QR avec un label optionnel en dessous.
    Retourne une chaîne LaTeX ou "" si pas de QR code.
    """
    if not prefs.qr_code_path:
        return ""
    label_tex = ""
    if prefs.qr_code_label:
        label_tex = f"\n\\\\[2pt]{{\\tiny\\color{{white}} {_latex_escape(prefs.qr_code_label)}}}"
    return (
        "\\vspace{0.3cm}\n"
        "\\hspace{0.2cm}\\safeimage{qrcode.png}{0.60\\linewidth}"
        + label_tex
    )


def _sidebar_photo_block(prefs: CVPreferences) -> str:
    """Photo d'identité ronde/carrée dans le bandeau gauche (optionnel).

    Différente de la photo header — permet d'avoir la photo dans le bandeau
    si l'utilisateur préfère ce placement.
    """
    if not prefs.sidebar_photo_path:
        return ""
    return (
        "\\begin{center}\n"
        "\\includegraphics[width=0.55\\linewidth,height=0.55\\linewidth,"
        "keepaspectratio]{sidebar_photo.png}\n"
        "\\end{center}\n"
        "\\vspace{0.2cm}"
    )


# ── Template loader ──────────────────────────────────────────────────────────
def _load_template(name: str) -> str:
    name = (name or "optimum").lower()
    if name not in {"optimum", "minimal"}:
        name = "optimum"
    path = TEMPLATES_DIR / f"{name}.tex"
    return path.read_text(encoding="utf-8")


def _inject_style(template: str, prefs: CVPreferences) -> str:
    """Remplace les placeholders globaux (couleurs, langue, photo, QR code)."""
    babel = "english" if prefs.language.lower().startswith("en") else "french"
    accent_rgb = _hex_to_rgb(prefs.accent_hex)
    leftbg_rgb = _hex_to_rgb(prefs.leftbg_hex)
    is_flat = "FLAT" in template
    photo_optimum = _photo_block(prefs)
    photo_minimal = _photo_block_minimal(prefs)
    qr_block = _qr_code_block(prefs)
    sidebar_photo = _sidebar_photo_block(prefs)
    # Headings — l'offre peut être en anglais aussi.
    if babel == "english":
        exp_heading = "Professional Experience"
        edu_heading = "Education"
    else:
        exp_heading = "Expériences professionnelles"
        edu_heading = "Formation"
    repl = {
        "{{BABEL_LANG}}": babel,
        "{{ACCENT_RGB}}": accent_rgb,
        "{{LEFTBG_RGB}}": leftbg_rgb,
        "{{PHOTO_BLOCK}}": photo_minimal if is_flat else photo_optimum,
        "{{PHOTO_ID_SIDEBAR_BLOCK}}": sidebar_photo,
        "{{QR_CODE_BLOCK}}": qr_block,
        "{{EXPERIENCE_HEADING}}": exp_heading,
        "{{EDUCATION_HEADING}}": edu_heading,
    }
    out = template
    for k, v in repl.items():
        out = out.replace(k, v)
    return out


# ── Constructeurs de blocs LaTeX ─────────────────────────────────────────────

def _build_section_block_optimum(title: str, items: list[str]) -> str:
    """Bloc colonne gauche (fond coloré, texte blanc)."""
    if not items:
        return ""
    rows = "\n".join(
        f"\\item \\textcolor{{white}}{{{_latex_escape(str(i))}}}" for i in items
    )
    return (
        f"{{\\color{{white}}\\sffamily\\bfseries {title}}}\n\n"
        f"\\vspace{{0.25cm}}\n\n"
        f"\\begin{{itemize}}[label=\\textcolor{{white}}{{$\\blacktriangleright$}}]\n"
        f"{rows}\n"
        f"\\end{{itemize}}\n\n"
        f"\\vspace{{0.55cm}}\n"
    )


def _build_section_block_minimal(title: str, items: list[str]) -> str:
    """Bloc section flat (template minimal, texte noir)."""
    if not items:
        return ""
    rows = "\n".join(f"\\item {_latex_escape(str(i))}" for i in items)
    return (
        f"\\section*{{{title}}}\n"
        f"\\begin{{itemize}}\n"
        f"{rows}\n"
        f"\\end{{itemize}}\n"
    )


def _build_experiences_block(
    experiences: list[dict], is_optimum: bool, extra_assets: Optional[dict] = None
) -> str:
    blocks = []
    for exp in experiences:
        title   = _latex_escape(exp.get("title", ""))
        company_raw = exp.get("company", "")
        company = _latex_escape(company_raw)
        dates   = _latex_escape(exp.get("dates", ""))
        bullets = exp.get("bullets", [])
        rows = "\n".join(f"\\item {_latex_escape(str(b))}" for b in bullets)

        # Logo de l'entreprise (depuis templates/assets/ ou extra_assets)
        logo = _find_logo(company_raw, extra_filenames=extra_assets)
        logo_tex = f"\\safeimage{{{logo}}}{{0.20\\textwidth}}" if logo else ""

        if is_optimum:
            block = (
                f"\\textbf{{{title} - {company}}}\n\n"
                f"\\emph{{{dates}}}\\hfill {logo_tex}\\\\\n\n"
                f"\\begin{{itemize}}[label=\\textcolor{{accent}}{{$\\blacktriangleright$}}]\n"
                f"{rows}\n"
                f"\\end{{itemize}}\n\n"
                f"\\vspace{{0.1cm}}\n"
            )
        else:
            block = (
                f"\\textbf{{{title} - {company}}} \\hfill \\emph{{{dates}}}\n"
                f"\\begin{{itemize}}\n"
                f"{rows}\n"
                f"\\end{{itemize}}\n\n"
            )
        blocks.append(block)
    return "\n".join(blocks)


def _build_education_block(education: list[dict], extra_assets: Optional[dict] = None) -> str:
    lines = []
    for edu in education:
        degree = _latex_escape(edu.get("degree", ""))
        school_raw = edu.get("school", "")
        school = _latex_escape(school_raw)
        dates  = _latex_escape(edu.get("dates", ""))

        # Logo de l'école (depuis templates/assets/ ou extra_assets)
        logo = _find_logo(school_raw, extra_filenames=extra_assets)
        logo_tex = f"\\safeimage{{{logo}}}{{0.10\\textwidth}}" if logo else ""

        lines.append(
            f"\\textbf{{{degree} - {school}}} \\hfill {logo_tex} \\emph{{{dates}}}\\\\[2pt]"
        )
    return ("\n".join(lines) + "\n\\vspace{0.2cm}") if lines else ""


def _build_latex_from_json(data: dict, template: str, prefs: CVPreferences) -> str:
    """Injecte les données JSON dans le template — la structure LaTeX ne change jamais."""
    is_optimum = "FLAT" not in template  # optimum a des sections colorées; minimal a des _FLAT
    is_fr = not prefs.language.lower().startswith("en")

    # ── Summary : liste de bullets ou string (rétro-compat) ─────────────────
    summary_raw = data.get("summary", "")
    if isinstance(summary_raw, list) and summary_raw:
        _bullets = "\n".join(
            f"\\item {_latex_escape(str(b))}" for b in summary_raw if str(b).strip()
        )
        summary_tex = (
            "\\begin{itemize}["
            "label=\\textcolor{accent}{$\\blacktriangleright$},"
            "noitemsep,topsep=2pt,leftmargin=*]\n"
            f"{_bullets}\n"
            "\\end{itemize}"
        )
    else:
        summary_tex = _latex_escape(str(summary_raw)) if summary_raw else ""

    # ── Champs simples ───────────────────────────────────────────────────────
    simple: dict[str, str] = {
        "{{FULL_NAME}}":        _latex_escape(data.get("full_name", "")),
        "{{HEADER_TITLE}}":     _latex_escape(data.get("header_title", "")),
        "{{EMAIL}}":            _latex_escape(data.get("email", "")),
        "{{LINKEDIN_URL}}":     data.get("linkedin_url", ""),
        "{{LINKEDIN_HANDLE}}":  _latex_escape(data.get("linkedin_handle", "")),
        "{{PHONE}}":            _latex_escape(data.get("phone", "")),
        "{{LOCATION}}":         _latex_escape(data.get("location", "")),
        "{{SUMMARY}}":          summary_tex,
    }
    result = template
    for k, v in simple.items():
        result = result.replace(k, v)

    # ── Labels de sections ───────────────────────────────────────────────────
    labels = {
        "skills":         "COMPÉTENCES"   if is_fr else "SKILLS",
        "tools":          "OUTILS"        if is_fr else "TOOLS",
        "languages":      "LANGUES"       if is_fr else "LANGUAGES",
        "certifications": "CERTIFICATIONS",
        "qualities":      "QUALITÉS"      if is_fr else "QUALITIES",
        "interests":      "INTÉRÊTS"      if is_fr else "INTERESTS",
        "awards":         "DISTINCTIONS"  if is_fr else "AWARDS",
    }

    # ── Sections colonne gauche (optimum) ────────────────────────────────────
    if is_optimum:
        for key, label in labels.items():
            ph = f"{{{{{key.upper()}_SECTION}}}}"
            items = data.get(key, [])
            result = result.replace(ph, _build_section_block_optimum(label, items))
    # ── Sections flat (minimal) ──────────────────────────────────────────────
    else:
        for key, label in labels.items():
            ph = f"{{{{{key.upper()}_SECTION_FLAT}}}}"
            items = data.get(key, [])
            result = result.replace(ph, _build_section_block_minimal(label, items))

    # ── Expériences & Formation ──────────────────────────────────────────────
    result = result.replace(
        "{{EXPERIENCES_BLOCK}}",
        _build_experiences_block(data.get("experiences", []), is_optimum, prefs.extra_assets),
    )
    result = result.replace(
        "{{EDUCATION_BLOCK}}",
        _build_education_block(data.get("education", []), prefs.extra_assets),
    )
    return result


def _try_fix_json(raw: str) -> str:
    """Corrections sur un JSON mal formé retourné par le LLM."""
    # 1. Extraire l'objet JSON s'il y a du texte autour
    m = re.search(r"\{.*\}", raw, re.DOTALL)
    raw = m.group(0) if m else raw
    # 2. Trailing commas
    raw = re.sub(r",(\s*[}\]])", r"\1", raw)
    # 3. Backslashes LaTeX invalides dans les strings JSON.
    #    JSON n'autorise que \" \\ \/ \b \f \n \r \t \uXXXX.
    #    On double-échappe tous les \ non-standards.
    _VALID = set('"\\bfnrtu/')
    out = []
    i = 0
    while i < len(raw):
        ch = raw[i]
        if ch == '\\' and i + 1 < len(raw):
            nxt = raw[i + 1]
            if nxt not in _VALID:
                out.append('\\\\')   # double-escaper pour rendre le JSON valide
            else:
                out.append('\\')
            i += 1
        else:
            out.append(ch)
        i += 1
    return ''.join(out)


# ── Limites colonne gauche (évite le débordement A4) ────────────────────────
# La colonne gauche (0.275\textwidth) tient ~35-38 items au total.
# Chaque section a un header (~1 ligne) + ses items (certains wrappent sur 2).
# Budget : 8+8+5+4+3+3+4 = 35 items + 7 headers → ~42 lignes visuelles, safe.
_LEFT_COL_LIMITS: dict[str, int] = {
    "skills":         8,
    "tools":          8,
    "qualities":      5,
    "interests":      3,
    "awards":         3,
    "certifications": 4,
    "languages":      5,
}


def _cap_left_column(data: dict) -> dict:
    """Tronque les listes de la colonne gauche pour éviter le débordement LaTeX."""
    for key, limit in _LEFT_COL_LIMITS.items():
        items = data.get(key)
        if isinstance(items, list) and len(items) > limit:
            data[key] = items[:limit]
    return data


# ── Génération CV ────────────────────────────────────────────────────────────
def _strip_code_fences(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n?", "", text)
        text = re.sub(r"\n?```\s*$", "", text)
    return text.strip()


def _strip_empty_sections(latex: str) -> str:
    """Filet de sécurité — supprime les sections dont l'itemize est vide.

    Reconnait deux patterns :
    1. Bandeau gauche optimum : `{\\color{white}\\sffamily\\bfseries TITRE} ...
       \\begin{itemize}...\\end{itemize}` où l'itemize ne contient pas `\\item`.
    2. `\\section*{...}` suivi d'un itemize sans `\\item`.
    """
    # Pattern 1 — colonne gauche optimum.
    pat1 = re.compile(
        r"\{\\color\{white\}\\sffamily\\bfseries[^}]+\}.*?"
        r"\\begin\{itemize\}\[[^\]]*\](?P<items>.*?)\\end\{itemize\}"
        r"(?:\s*\\vspace\{[^}]+\})?",
        re.DOTALL,
    )
    def _maybe1(m: re.Match) -> str:
        items = m.group("items")
        if "\\item" in items:
            return m.group(0)
        return ""
    latex = pat1.sub(_maybe1, latex)
    # Pattern 2 — section minimal.
    pat2 = re.compile(
        r"\\section\*\{[^}]+\}\s*\\begin\{itemize\}(?P<items>.*?)\\end\{itemize\}",
        re.DOTALL,
    )
    latex = pat2.sub(
        lambda m: "" if "\\item" not in m.group("items") else m.group(0),
        latex,
    )
    # Pattern 3 — nettoyer les triples sauts.
    latex = re.sub(r"\n{3,}", "\n\n", latex)
    return latex


def generate_optimized_cv(
    job_offer: str,
    source_cv: str,
    prefs: Optional[CVPreferences] = None,
) -> str:
    """Génère un CV LaTeX en préservant EXACTEMENT la structure du template.

    Le LLM ne renvoie que les données (JSON).
    On injecte nous-mêmes dans le template — structure toujours identique.
    """
    prefs = prefs or CVPreferences()

    # Charger le template + injecter style (couleurs, babel, photo)
    raw_template = _load_template(prefs.template)
    template = _inject_style(raw_template, prefs)

    # Injecter les headings traduits
    is_en = prefs.language.lower().startswith("en")
    template = template.replace(
        "{{EXPERIENCE_HEADING}}",
        "Professional Experience" if is_en else "Expériences professionnelles",
    )
    template = template.replace(
        "{{EDUCATION_HEADING}}",
        "Education" if is_en else "Formation",
    )

    language_label = "Anglais" if is_en else "Français"

    prompt = _PROMPT_STRUCTURED_JSON.format(
        language_label=language_label,
        job_offer=(job_offer or "").strip()[:8000],
        source_cv=(source_cv or "").strip()[:8000],
    )

    raw = generate_with_llm(prompt, temperature=0.3)
    raw = _strip_code_fences(raw)

    # Parser le JSON
    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        raw_fixed = _try_fix_json(raw)
        try:
            data = json.loads(raw_fixed)
        except Exception as e:
            raise RuntimeError(
                f"Le LLM n'a pas retourné un JSON valide : {e}\n---\n{raw[:800]}"
            )

    # Limiter la colonne gauche pour éviter le débordement A4
    data = _cap_left_column(data)

    # Construire le LaTeX en injectant les données dans le template fixe
    latex = _build_latex_from_json(data, template, prefs)

    if r"\documentclass" not in latex:
        raise RuntimeError("Injection échouée : \\documentclass absent du LaTeX final.")

    return latex


# ── Génération lettre ────────────────────────────────────────────────────────
def generate_cover_letter(
    job_offer: str,
    optimized_cv: str,
    prefs: Optional[CVPreferences] = None,
) -> str:
    prefs = prefs or CVPreferences()
    language_label = "Anglais" if prefs.language.lower().startswith("en") else "Français"
    prompt = _PROMPT_LETTER.format(
        language_label=language_label,
        job_offer=(job_offer or "").strip()[:8000],
        optimized_cv=(optimized_cv or "").strip()[:8000],
    )
    return _strip_code_fences(generate_with_llm(prompt, temperature=0.55))


# ── Extraction nom / titre / entreprise ──────────────────────────────────────
def _guess_name_from_cv_latex(cv_latex: str) -> str:
    m = re.search(r"\\textbf\{\s*([^{}\\]+?)\s*\}\s*\}", cv_latex)
    return m.group(1).strip() if m else ""


def _guess_title_from_cv_latex(cv_latex: str) -> str:
    matches = re.findall(r"\\textbf\{\s*([^{}\\]+?)\s*\}\s*\}", cv_latex)
    return matches[1].strip() if len(matches) > 1 else ""


def _guess_company(job_offer: str) -> str:
    """Heuristique : on cherche le 1er nom propre majuscule plausible."""
    txt = job_offer or ""
    # Pattern 1 — "chez X" / "at X" / "rejoindre X"
    m = re.search(
        r"(?:chez|at|pour|rejoindre|au sein de)\s+"
        r"([A-Z][\wÀ-ÿ&\.\- ]{2,40})",
        txt,
    )
    if m:
        return m.group(1).strip().rstrip(",.")
    # Pattern 2 — ligne de l'offre qui commence par un mot capitalisé long
    for line in txt.splitlines():
        line = line.strip()
        if 3 < len(line) < 50 and line[0].isupper() and " " in line:
            if any(w in line.lower() for w in ("ltd", "sa", "sas", "sarl", "concession", "airports", "group", "vinci")):
                return line
    return ""


# ── Lettre LaTeX ─────────────────────────────────────────────────────────────
def _build_letter_latex(
    body: str,
    prefs: CVPreferences,
    *,
    sender_name: str,
    job_title: str,
    company: str,
) -> str:
    template = TEMPLATE_LETTER.read_text(encoding="utf-8")
    is_en = prefs.language.lower().startswith("en")
    salutation = "Dear Hiring Manager" if is_en else "Madame, Monsieur"
    closing = (
        "Yours sincerely,"
        if is_en
        else "Je vous prie d'agréer, Madame, Monsieur, l'expression de mes salutations distinguées."
    )
    subject_label = "Application for" if is_en else "Candidature au poste de"
    subject = f"{subject_label} {_latex_escape(job_title or '[poste]')}"
    if company:
        subject += " " + ("at " if is_en else "chez ") + _latex_escape(company)
    repl = {
        "{{SENDER_NAME}}": _latex_escape(sender_name or "[Votre nom]"),
        "{{SENDER_ADDRESS_BLOCK}}": "",
        "{{SENDER_EMAIL}}": "",
        "{{SENDER_PHONE}}": "",
        "{{SENDER_CITY}}": _latex_escape("Paris"),
        "{{RECIPIENT_COMPANY}}": _latex_escape(company or ""),
        "{{RECIPIENT_ADDRESS_BLOCK}}": "",
        "{{DATE}}": _english_date() if is_en else _french_date(),
        "{{SUBJECT}}": subject,
        "{{SALUTATION}}": salutation,
        "{{LETTER_BODY}}": _paragraphs_to_latex(body),
        "{{CLOSING}}": closing,
    }
    for k, v in repl.items():
        template = template.replace(k, v)
    return template


# ── Compilation PDF ──────────────────────────────────────────────────────────
def _tighten_latex_for_one_page(latex_src: str, level: int = 1) -> str:
    """Réduit marges + police + interligne pour forcer 1 page.

    level=1 : resserrement doux   (9.5pt→9pt,  margins 1.0cm)
    level=2 : resserrement fort   (9.5pt→8.5pt, margins 0.75cm)
    """
    if level == 1:
        font = "9pt"
        margin = r"\usepackage[margin=1.0cm,top=0.8cm,bottom=0.7cm]{geometry}"
        linespread = r"\linespread{0.92}\selectfont"
        sidebar_w = r"\setlength{\sidebarw}{0.403\paperwidth}"
        box_height = r"\begin{minipage}[t][21cm][t]"
        enlarge = r"\enlargethispage*{2\baselineskip}"
        vspace_map = [
            (r'\vspace{0.8em}',   r'\vspace{0.4em}'),
            (r'\vspace{0.6em}',   r'\vspace{0.3em}'),
            (r'\vspace{0.45cm}',  r'\vspace{0.2cm}'),
            (r'\vspace{0.3cm}',   r'\vspace{0.15cm}'),
            (r'\vspace{0.2cm}',   r'\vspace{0.1cm}'),
            (r'\vspace{0.15cm}',  r'\vspace{0.08cm}'),
            (r'\vspace{0.1cm}',   r'\vspace{0.05cm}'),
            (r'\\[4pt]',          r'\\[2pt]'),
            (r'\\[6pt]',          r'\\[3pt]'),
        ]
        list_inject = (
            r'\setlist{nosep,topsep=2pt,parsep=0pt,partopsep=0pt,leftmargin=1.2em}' + '\n'
        )
    else:
        font = "8.5pt"
        margin = r"\usepackage[margin=0.75cm,top=0.55cm,bottom=0.55cm]{geometry}"
        linespread = r"\linespread{0.88}\selectfont"
        sidebar_w = r"\setlength{\sidebarw}{0.400\paperwidth}"
        box_height = r"\begin{minipage}[t][22cm][t]"
        enlarge = r"\enlargethispage*{4\baselineskip}"
        vspace_map = [
            (r'\vspace{0.8em}',   r'\vspace{0.05em}'),
            (r'\vspace{0.6em}',   r'\vspace{0.03em}'),
            (r'\vspace{0.5em}',   r'\vspace{0.02em}'),
            (r'\vspace{0.45cm}',  r'\vspace{0.1cm}'),
            (r'\vspace{0.3cm}',   r'\vspace{0.08cm}'),
            (r'\vspace{0.2cm}',   r'\vspace{0.04cm}'),
            (r'\vspace{0.15cm}',  r'\vspace{0.03cm}'),
            (r'\vspace{0.1cm}',   r'\vspace{0.02cm}'),
            (r'\vspace{0.05cm}',  r'\vspace{0.01cm}'),
            (r'\\[4pt]',          r'\\[1pt]'),
            (r'\\[6pt]',          r'\\[2pt]'),
            (r'\\[2pt]',          r'\\[0pt]'),
        ]
        list_inject = (
            r'\setlist{nosep,topsep=1pt,parsep=0pt,partopsep=0pt,leftmargin=1.2em}' + '\n'
        )

    # Réduire la taille de police
    patched = re.sub(r'(\\documentclass\[)[\d.]+pt', r'\g<1>' + font, latex_src)
    # Réduire les marges (geometry) — lambda évite les escapes regex dans la repl
    _margin = margin
    patched = re.sub(r'\\usepackage\[[^\]]*\]\{geometry\}', lambda m: _margin, patched)
    # Mettre à jour la largeur du bandeau (tikz — ignoré si absent)
    _sidebar_w = sidebar_w
    patched = re.sub(
        r'\\setlength\{\\sidebarw\}\{[^}]+\}',
        lambda m: _sidebar_w,
        patched,
    )
    # Ajuster la hauteur du minipage bandeau (colorbox approach)
    _box_h = box_height
    patched = re.sub(
        r'\\begin\{minipage\}\[t\]\[[\d.]+cm\]\[t\]',
        lambda m: _box_h,
        patched,
    )
    # Réduire les espacements
    for old, new in vspace_map:
        patched = patched.replace(old, new)
    # Injecter interligne + itemize compact avant \begin{document}
    inject_pre = linespread + '\n' + list_inject
    patched = patched.replace(r'\begin{document}', inject_pre + r'\begin{document}')
    # \enlargethispage juste après \begin{document}
    after_begin = r'\begin{document}' + '\n'
    patched = patched.replace(after_begin, after_begin + enlarge + '\n', 1)
    return patched


def _count_pdf_pages(pdf_path: Path) -> int:
    """Retourne le nombre de pages du PDF via PyMuPDF."""
    try:
        import fitz
        doc = fitz.open(str(pdf_path))
        n = len(doc)
        doc.close()
        return n
    except Exception:
        return 1  # en cas d'erreur, on suppose 1 page


def _compile_latex(
    latex_src: str,
    base_name: str,
    extra_assets: Optional[dict] = None,
    enforce_one_page: bool = True,
    prefs: Optional[CVPreferences] = None,
) -> tuple[Optional[bytes], list[str]]:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    safe = re.sub(r"[^A-Za-z0-9_\-]+", "_", base_name) or "document"
    tex_path = OUTPUT_DIR / f"{safe}.tex"
    pdf_path = OUTPUT_DIR / f"{safe}.pdf"

    # Copier les assets (logos, photo, QR code) dans outputs/ pour que pdflatex les trouve
    _copy_assets(OUTPUT_DIR, extra_assets, prefs=prefs)

    errors: list[str] = []
    if shutil.which("pdflatex") is None:
        return None, ["pdflatex introuvable, installez TeX Live ou MiKTeX."]

    def _run_pdflatex(src: str) -> tuple[bool, str]:
        tex_path.write_text(src, encoding="utf-8")
        cmd = [
            "pdflatex",
            "-interaction=nonstopmode",
            "-output-directory", str(OUTPUT_DIR),
            str(tex_path),
        ]
        for run in (1, 2):
            try:
                res = subprocess.run(
                    cmd, capture_output=True, timeout=300, cwd=str(OUTPUT_DIR)
                )
            except subprocess.TimeoutExpired:
                return False, "pdflatex timeout (300s)"
            except Exception as e:
                return False, f"pdflatex erreur : {e}"
            raw = res.stdout or b""
            if isinstance(raw, bytes):
                stdout = raw.decode("utf-8", errors="replace")
                if "â€" in stdout or "â" in stdout:
                    stdout = raw.decode("cp1252", errors="replace")
            else:
                stdout = raw
            if res.returncode != 0 and run == 2 and not pdf_path.exists():
                tail = "\n".join(stdout.splitlines()[-30:])
                return False, f"pdflatex exit={res.returncode}.\n{tail}"
        return True, ""

    # Passe 1 — compilation normale
    ok, err = _run_pdflatex(latex_src)
    if not ok:
        errors.append(err)
        return None, errors
    if not pdf_path.exists():
        return None, errors or ["PDF non produit."]

    # Passe 2 — si > 1 page, resserrement en 2 niveaux
    if enforce_one_page:
        pages = _count_pdf_pages(pdf_path)
        if pages > 1:
            logger.info("CV %d pages — resserrement niveau 1.", pages)
            tighter1 = _tighten_latex_for_one_page(latex_src, level=1)
            ok2, err2 = _run_pdflatex(tighter1)
            if ok2 and pdf_path.exists():
                pages2 = _count_pdf_pages(pdf_path)
                if pages2 > 1:
                    logger.info("Toujours %d pages — resserrement niveau 2.", pages2)
                    tighter2 = _tighten_latex_for_one_page(latex_src, level=2)
                    ok3, err3 = _run_pdflatex(tighter2)
                    if ok3 and pdf_path.exists():
                        pages3 = _count_pdf_pages(pdf_path)
                        if pages3 > 1:
                            logger.warning("CV %d pages après niveau 2 — contenu trop dense.", pages3)
                            errors.append(
                                f"⚠ CV généré en {pages3} pages malgré le resserrement maximum. "
                                "Conseil : réduisez le nombre de bullets ou retirez une expérience."
                            )
                    else:
                        logger.warning("Niveau 2 échoué (%s) — PDF niveau 1 conservé.", err3)
                        _run_pdflatex(tighter1)
            else:
                logger.warning("Niveau 1 échoué (%s) — PDF original conservé.", err2)
                _run_pdflatex(latex_src)

    if not pdf_path.exists():
        return None, errors or ["PDF non produit."]
    return pdf_path.read_bytes(), errors


def create_pdf(latex_or_text: str, output_filename: str) -> bytes:
    base = Path(output_filename).stem or "document"
    pdf_bytes, errors = _compile_latex(latex_or_text, base)
    if pdf_bytes is None:
        raise RuntimeError(
            f"Échec compilation {output_filename} : {'; '.join(errors) or 'inconnu'}"
        )
    return pdf_bytes


# ── Naming ───────────────────────────────────────────────────────────────────
def _slug(s: str, max_len: int = 40) -> str:
    if not s:
        return ""
    nfkd = unicodedata.normalize("NFKD", s)
    ascii_str = "".join(c for c in nfkd if not unicodedata.combining(c))
    out = []
    for ch in ascii_str:
        if ch.isalnum():
            out.append(ch)
        elif ch in " -_'":
            out.append("_")
    s = "".join(out)
    while "__" in s:
        s = s.replace("__", "_")
    return s.strip("_")[:max_len]


def build_output_basename(candidate: str, job_title: str, company: str) -> str:
    """`Nom_Poste_Entreprise` slug, fallback `CV`/`LM`."""
    parts = [_slug(candidate), _slug(job_title), _slug(company)]
    parts = [p for p in parts if p]
    return "_".join(parts) if parts else ""


# ── Pipeline complet ─────────────────────────────────────────────────────────
def run_optimum_pipeline(
    job_offer: str,
    source_cv_text: str,
    prefs: Optional[CVPreferences] = None,
) -> dict:
    prefs = prefs or CVPreferences()
    out: dict = {
        "cv_latex": "", "cv_pdf_bytes": None, "cv_errors": [],
        "letter_body": "", "letter_latex": "",
        "letter_pdf_bytes": None, "letter_errors": [],
        "candidate_name": "", "job_title": "", "company": prefs.company,
        "cv_filename": "CV.pdf", "letter_filename": "Lettre_Motivation.pdf",
    }

    # 1. CV
    cv_latex = generate_optimized_cv(job_offer, source_cv_text, prefs)
    out["cv_latex"] = cv_latex
    candidate = _guess_name_from_cv_latex(cv_latex)
    job_title = _guess_title_from_cv_latex(cv_latex)
    company = prefs.company.strip() or _guess_company(job_offer)
    out["candidate_name"] = candidate
    out["job_title"] = job_title
    out["company"] = company

    base = build_output_basename(candidate, job_title, company)
    cv_base = f"{base}_CV" if base else "CV"
    lm_base = f"{base}_LM" if base else "Lettre_Motivation"
    out["cv_filename"] = f"{cv_base}.pdf"
    out["letter_filename"] = f"{lm_base}.pdf"

    cv_bytes, cv_errors = _compile_latex(
        cv_latex, cv_base,
        extra_assets=prefs.extra_assets,
        enforce_one_page=True,
        prefs=prefs,
    )
    out["cv_pdf_bytes"] = cv_bytes
    out["cv_errors"] = cv_errors

    # 2. Lettre — seulement si demandée
    if prefs.generate_letter:
        body = generate_cover_letter(job_offer, cv_latex, prefs)
        out["letter_body"] = body
        letter_latex = _build_letter_latex(
            body, prefs,
            sender_name=candidate,
            job_title=job_title,
            company=company,
        )
        out["letter_latex"] = letter_latex
        lm_bytes, lm_errors = _compile_latex(letter_latex, lm_base, enforce_one_page=False)
        out["letter_pdf_bytes"] = lm_bytes
        out["letter_errors"] = lm_errors

    return out
