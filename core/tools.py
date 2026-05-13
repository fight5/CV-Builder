"""Utility tools for text extraction, LaTeX formatting, and keyword analysis."""

import re
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_pdf(path: str) -> str:
    """Extract plain text from a PDF file using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text("text"))
        doc.close()
        return "\n".join(text_parts)
    except ImportError:
        logger.error("PyMuPDF (fitz) is not installed. Run: pip install pymupdf")
        raise
    except Exception as e:
        logger.error(f"Failed to extract text from PDF {path}: {e}")
        raise


def extract_text_from_docx(path: str) -> str:
    """Extract plain text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        # Also extract table content
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except ImportError:
        logger.error("python-docx is not installed. Run: pip install python-docx")
        raise
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX {path}: {e}")
        raise


def compute_keyword_density(text: str, keywords: list) -> float:
    """Return the fraction of keywords found in the given text (0.0 to 1.0)."""
    if not keywords:
        return 0.0
    text_lower = text.lower()
    found = 0
    for kw in keywords:
        if kw.lower() in text_lower:
            found += 1
    return round(found / len(keywords), 4)


def latex_escape(text: str) -> str:
    """Escape special LaTeX characters in a string."""
    if not isinstance(text, str):
        text = str(text)
    # Order matters — backslash must be first
    replacements = [
        ("\\", r"\textbackslash{}"),
        ("&", r"\&"),
        ("%", r"\%"),
        ("$", r"\$"),
        ("#", r"\#"),
        ("_", r"\_"),
        ("{", r"\{"),
        ("}", r"\}"),
        ("~", r"\textasciitilde{}"),
        ("^", r"\textasciicircum{}"),
        ("<", r"\textless{}"),
        (">", r"\textgreater{}"),
    ]
    for char, escaped in replacements:
        text = text.replace(char, escaped)
    return text


def format_experience_latex(exp_dict: dict) -> str:
    """Return a LaTeX snippet for a single work experience entry."""
    title = latex_escape(exp_dict.get("title", ""))
    company = latex_escape(exp_dict.get("company", ""))
    dates = latex_escape(exp_dict.get("dates", ""))
    location = latex_escape(exp_dict.get("location", ""))
    description = exp_dict.get("description", "")
    achievements = exp_dict.get("achievements", [])

    lines = [
        r"\experienceentry{" + title + r"}{" + company + r"}{" + dates + r"}{" + location + r"}{",
        r"  \begin{itemize}[leftmargin=*, topsep=2pt, itemsep=1pt]",
    ]

    if description and description.strip():
        lines.append(r"    \item " + latex_escape(description))

    for ach in achievements:
        if ach and ach.strip():
            lines.append(r"    \item " + latex_escape(ach))

    lines.append(r"  \end{itemize}")
    lines.append(r"}")
    return "\n".join(lines)


def sanitize_for_ats(text: str) -> str:
    """Remove characters that may confuse ATS parsers."""
    # Remove zero-width spaces, non-breaking spaces, soft hyphens
    text = re.sub(r"[​­﻿]", "", text)
    # Replace curly quotes with straight quotes
    text = text.replace("‘", "'").replace("’", "'")
    text = text.replace("“", '"').replace("”", '"')
    # Replace en/em dashes with regular dash
    text = text.replace("–", "-").replace("—", "-")
    return text


def extract_email(text: str) -> Optional[str]:
    """Extract first email address found in text."""
    match = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else None


def extract_phone(text: str) -> Optional[str]:
    """Extract first phone number found in text."""
    match = re.search(r"(?:\+?\d{1,3}[\s\-.]?)?\(?\d{2,4}\)?[\s\-.]?\d{2,4}[\s\-.]?\d{2,9}", text)
    return match.group(0).strip() if match else None


def normalize_date(date_str: str) -> str:
    """Normalize date strings to a consistent format for LaTeX output."""
    if not date_str:
        return ""
    # Common patterns: Jan 2022, January 2022, 01/2022, 2022-01
    date_str = date_str.strip()
    month_map = {
        "january": "Jan", "february": "Feb", "march": "Mar", "april": "Apr",
        "may": "May", "june": "Jun", "july": "Jul", "august": "Aug",
        "september": "Sep", "october": "Oct", "november": "Nov", "december": "Dec",
        "janvier": "Jan", "février": "Fév", "mars": "Mar", "avril": "Avr",
        "mai": "Mai", "juin": "Jun", "juillet": "Jul", "août": "Aoû",
        "septembre": "Sep", "octobre": "Oct", "novembre": "Nov", "décembre": "Déc",
    }
    lower = date_str.lower()
    for full, abbrev in month_map.items():
        lower = lower.replace(full, abbrev)
    return lower.strip()
